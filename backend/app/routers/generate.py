import os
import re
from datetime import datetime
from fastapi import APIRouter, HTTPException, UploadFile, File
from ..schemas.code_gen import CodeGenerationRequest, CodeGenerationResponse
from ..config.settings import get_db, MONGODB_URI
from ..services.db import get_or_create_quota, resolve_domains_if_needed
from ..generator import generator

router = APIRouter(tags=["Code Generation"])

PRICING = {
    "gpt-4o": {"input": 0.000005, "output": 0.000015},
    "gemini-3.5-flash": {"input": 0.000000075, "output": 0.0000003},
    "mistral": {"input": 0.000002, "output": 0.000006},
    "llama": {"input": 0.0000007, "output": 0.0000007},
    "kimi": {"input": 0.0000007, "output": 0.0000007}
}

def detect_tables_and_columns(sql_code: str, db_inst) -> tuple:
    cursor = db_inst["tableStatusNew"].find({"approvalStatus": "approved"})
    detected_tables = []
    detected_columns = []
    
    code_str = sql_code.strip()
    if "```" in code_str:
        blocks = re.findall(r'```(?:\w+)?\n(.*?)\n```', code_str, re.DOTALL)
        if blocks:
            code_str = blocks[0].strip()
        else:
            code_str = re.sub(r'```(?:\w+)?', '', code_str).strip()
            
    for doc in cursor:
        table_name = doc["tableName"]
        pattern = rf"\b{re.escape(table_name)}\b"
        if re.search(pattern, code_str, re.IGNORECASE):
            detected_tables.append(table_name)
            
            meta_store_doc = db_inst["semanticMetaStore"].find_one({"collection_name": table_name})
            columns = []
            if meta_store_doc and "fields" in meta_store_doc:
                columns = [f["field_name"] for f in meta_store_doc["fields"]]
            else:
                sample_doc = db_inst[table_name].find_one()
                if sample_doc:
                    columns = [key for key in sample_doc.keys() if key != '_id']
                    
            for col in columns:
                col_pattern = rf"\b{re.escape(col)}\b"
                if re.search(col_pattern, code_str, re.IGNORECASE):
                    detected_columns.append(col)
                    
    return list(set(detected_tables)), list(set(detected_columns))

@router.post("/generate/estimate")
async def estimate_tokens(request: CodeGenerationRequest):
    try:
        db = get_db()
        
        # Apply Input Guardrails
        from ..guardrails import run_input_guardrails
        request.logic = await run_input_guardrails(
            userId=request.userId,
            role=request.role,
            logic_prompt=request.logic,
            tables=request.tables,
            columns=request.columns,
            is_conversion=getattr(request, "is_conversion", False),
            model=request.model
        )
        
        # Dynamically decide model via supervisor
        from ..agents.supervisor import supervisor_decide_models
        decisions = supervisor_decide_models(request.logic, request.format, request.tables, request.model)
        model_key = decisions["code_generation"]["model"]
            
        # Check quota status before doing the estimate
        if request.role:
            await resolve_domains_if_needed(request, db)
            quota = get_or_create_quota(db, request.role)
            
            # Check model tokens limit
            model_quota = quota.get("limits", {}).get(model_key, {"total_tokens": 1000000, "used_tokens": 0})
            if model_quota.get("used_tokens", 0) >= model_quota.get("total_tokens", 1000000):
                raise HTTPException(
                    status_code=429, 
                    detail=f"Quota Exceeded: You have run out of tokens/credits for {model_key} model under the {request.role} role."
                )
            # Check financial budget balance
            if quota.get("remaining_balance_usd", 15.00) <= 0:
                raise HTTPException(
                    status_code=429, 
                    detail=f"Quota Exceeded: The credit balance for the {request.role} role is fully depleted ($0.00 remaining)."
                )

        schema_context = ""
        try:
            await resolve_domains_if_needed(request, db)
            schema_docs = list(db['semanticMetaStore'].find({"collection_name": {"$in": request.tables}}))
            
            schema_context_list = []
            for doc in schema_docs:
                col_name = doc.get("collection_name")
                desc = doc.get("description", "")
                pk = doc.get("primary_key", "")
                fields = doc.get("fields", [])
                
                fields_desc = []
                for f in fields:
                    f_name = f.get("field_name")
                    f_type = f.get("data_type")
                    f_desc = f.get("description")
                    fields_desc.append(f"- {f_name} ({f_type}): {f_desc}")
                    
                relations = doc.get("relations", [])
                relations_desc = []
                for r in relations:
                    relations_desc.append(f"Foreign key `{r.get('local_field')}` links to `{r.get('referenced_collection')}({r.get('referenced_field')})`")
                    
                schema_info = f"Table: {col_name}\nDescription: {desc}\nPrimary Key: {pk}\nColumns:\n" + "\n".join(fields_desc)
                if relations_desc:
                    schema_info += "\nRelations:\n" + "\n".join(relations_desc)
                schema_context_list.append(schema_info)
            
            if schema_context_list:
                schema_context = "\n\n=== Table Schemas ===\n" + "\n\n".join(schema_context_list)
        except Exception as e:
            print(f"Failed to fetch schemas for estimation: {e}")

        prompt = f"""
        You are an expert Data Engineer and AI Assistant specializing in SDLC automation.
        Generate the requested code based on the following input:
        
        Format: {request.format}
        Tables: {", ".join(request.tables)}
        Columns: {", ".join(request.columns)}
        Logic: {request.logic}
        Sample Data Size: {request.sample_data_size}
        {schema_context}
        """
        approx_prompt_tokens = len(prompt) // 4
        approx_prompt_tokens += 300  # System/instructions padding
        approx_completion_tokens = 450
        
        # Use the dynamic model_key decided by the supervisor at start of method
        # model_key is already assigned at the top
        rates = PRICING.get(model_key, {"input": 0.000005, "output": 0.000015})
        cost = (approx_prompt_tokens * rates["input"]) + (approx_completion_tokens * rates["output"])
        
        return {
            "model": model_key,
            "approx_input_tokens": approx_prompt_tokens,
            "approx_output_tokens": approx_completion_tokens,
            "approx_cost_usd": round(cost, 6)
        }
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/generate", response_model=CodeGenerationResponse)
async def generate_code(request: CodeGenerationRequest):
    try:
        db = get_db()
        
        # Apply Input Guardrails
        from ..guardrails import run_input_guardrails
        request.logic = await run_input_guardrails(
            userId=request.userId,
            role=request.role,
            logic_prompt=request.logic,
            tables=request.tables,
            columns=request.columns,
            is_conversion=getattr(request, "is_conversion", False),
            model=request.model
        )
        
        role = request.role
        
        # Check Semantic Cache
        from ..services.embeddings import get_embedding, calculate_similarity
        from ..services.simulation_runner import run_simulation_logic
        from ..services.dq_profiler import calculate_dataframe_dq
        
        req_emb = await get_embedding(request.logic)
        
        cache_cursor = db["Semantic_Cache"].find({"format": request.format})
        best_match = None
        highest_similarity = 0.0
        
        for cache_entry in cache_cursor:
            entry_emb = cache_entry.get("embedding")
            similarity = calculate_similarity(request.logic, cache_entry["query"], req_emb, entry_emb)
            if similarity > highest_similarity:
                highest_similarity = similarity
                best_match = cache_entry
                
        if best_match and highest_similarity >= 0.90:
            # Increment and check hit count
            today_str = datetime.now().strftime("%Y-%m-%d")
            query_key = best_match["query"]
            
            db["Semantic_Cache_Hits"].update_one(
                {"query": query_key, "date": today_str},
                {"$inc": {"count": 1}},
                upsert=True
            )
            
            hit_doc = db["Semantic_Cache_Hits"].find_one({"query": query_key, "date": today_str})
            hit_count = hit_doc["count"] if hit_doc else 1
            
            if hit_count >= 3:
                # Return cached code, bypass LLM
                cached_code = best_match["code"]
                
                # Execute native simulation runner on cached code to get dynamic DQ Insights
                try:
                    sim_res = await run_simulation_logic(
                        tables=request.tables,
                        columns=request.columns,
                        generated_code=cached_code,
                        format_str=request.format,
                        sample_data_size=request.sample_data_size,
                        logic=request.logic,
                        role=request.role,
                        userId=request.userId,
                        mock_inputs=None
                    )
                    final_df = sim_res["final_dataframe"]
                    col_details = sim_res["column_details"]
                    dq = calculate_dataframe_dq(final_df, col_details)
                except Exception as e:
                    print(f"[ERROR] Simulation for cache hit failed: {e}")
                    dq = {
                        "row_count": 0, "null_values": 0, "duplicate_rows": 0,
                        "minimum": None, "maximum": None, "average": None,
                        "distinct_values": 0, "empty_strings": 0
                    }
                
                from ..schemas.code_gen import DQInsights
                result = CodeGenerationResponse(
                    generated_code=cached_code,
                    dq_insights=DQInsights(**dq),
                    flow_explanation=f"Served from Semantic Cache (Bypassed LLM due to {hit_count} hits today).\n\n=== LLM Configuration Details ===\nServed from Cache (Similarity: {highest_similarity:.2f})",
                    prompt_tokens=0,
                    completion_tokens=0,
                    insights=[],
                    personas=[]
                )
                
                detected_tables, detected_columns = detect_tables_and_columns(result.generated_code, db)
                result.detected_tables = detected_tables if detected_tables else request.tables
                result.detected_columns = detected_columns if detected_columns else request.columns
                return result
        
        # Dynamically decide model via supervisor
        from ..agents.supervisor import supervisor_decide_models
        decisions = supervisor_decide_models(request.logic, request.format, request.tables, request.model)
        model_key = decisions["code_generation"]["model"]
        request.model = model_key
        
        await resolve_domains_if_needed(request, db)
            
        if role:
            quota = get_or_create_quota(db, role)
            
            # Check model tokens limit
            model_quota = quota.get("limits", {}).get(model_key, {"total_tokens": 1000000, "used_tokens": 0})
            if model_quota.get("used_tokens", 0) >= model_quota.get("total_tokens", 1000000):
                raise HTTPException(
                    status_code=429, 
                    detail=f"Quota Exceeded: You have run out of tokens/credits for {model_key} model under the {role} role."
                )
            # Check financial budget balance
            if quota.get("remaining_balance_usd", 15.00) <= 0:
                raise HTTPException(
                    status_code=429, 
                    detail=f"Quota Exceeded: The credit balance for the {role} role is fully depleted ($0.00 remaining)."
                )

        from ..agents.supervisor import run_agent_workflow
        # Execute LangGraph Multi-Agent Workflow
        agent_result = await run_agent_workflow(request)
        
        from ..schemas.code_gen import DQInsights
        result = CodeGenerationResponse(
            generated_code=agent_result.get("generated_code", ""),
            dq_insights=DQInsights(**agent_result.get("dq_insights", {
                "row_count": 0, "null_values": 0, "duplicate_rows": 0, "distinct_values": 0, "empty_strings": 0
            })),
            flow_explanation=agent_result.get("flow_explanation", ""),
            prompt_tokens=agent_result.get("prompt_tokens", 0),
            completion_tokens=agent_result.get("completion_tokens", 0),
            insights=agent_result.get("insights", []),
            personas=agent_result.get("personas", [])
        )
        
        detected_tables, detected_columns = detect_tables_and_columns(result.generated_code, db)
        result.detected_tables = detected_tables if detected_tables else request.tables
        result.detected_columns = detected_columns if detected_columns else request.columns

        if role and result:
            p_tokens = result.prompt_tokens or 0
            c_tokens = result.completion_tokens or 0
            total_tokens = p_tokens + c_tokens
            
            rates = PRICING.get(model_key, {"input": 0.000005, "output": 0.000015})
            cost = (p_tokens * rates["input"]) + (c_tokens * rates["output"])
            
            limits = quota.get("limits", {})
            if model_key not in limits:
                limits[model_key] = {"total_tokens": 1000000, "used_tokens": 0}
            limits[model_key]["used_tokens"] = limits[model_key].get("used_tokens", 0) + total_tokens
            
            # Update database
            db["modelQuotas"].update_one(
                {"role": role},
                {
                    "$set": {
                        "limits": limits
                    },
                    "$inc": {
                        "remaining_balance_usd": -cost
                    }
                }
            )
            
            # Insert usage audit record
            log_doc = {
                "userId": request.userId or "unknown",
                "role": role,
                "timestamp": datetime.now(),
                "tokens_consumed": total_tokens,
                "cost": round(cost, 6)
            }
            db["roleTokenConsumption"].insert_one(log_doc)
            
        return result
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/generate/validate-input")
async def validate_input_guardrails_endpoint(request: CodeGenerationRequest):
    try:
        db = get_db()
        await resolve_domains_if_needed(request, db)
        
        checked = []
        
        # 1. Prompt Injection Prevention
        passed_pi = True
        error_pi = None
        try:
            from ..guardrails.injection import validate_prompt_injection
            validate_prompt_injection(request.logic)
            checked.append({"name": "Prompt Injection Prevention", "status": "Passed"})
        except HTTPException as he:
            passed_pi = False
            error_pi = he.detail
            checked.append({"name": "Prompt Injection Prevention", "status": "Failed", "message": error_pi})
            
        # 1b. Authentication Bypass Prevention (Always True Escape)
        passed_bypass = True
        error_bypass = None
        try:
            from ..guardrails.injection import validate_always_true_escape
            validate_always_true_escape(request.logic)
            checked.append({"name": "Authentication Bypass Prevention", "status": "Passed"})
        except HTTPException as he:
            passed_bypass = False
            error_bypass = he.detail
            checked.append({"name": "Authentication Bypass Prevention", "status": "Failed", "message": error_bypass})
            
        # 1c. Profanity & Content Moderation Check
        passed_profanity = True
        error_profanity = None
        try:
            from ..guardrails.profanity import validate_profanity
            await validate_profanity(request.logic)
            checked.append({"name": "Profanity & Moderation", "status": "Passed"})
        except HTTPException as he:
            passed_profanity = False
            error_profanity = he.detail
            checked.append({"name": "Profanity & Moderation", "status": "Failed", "message": error_profanity})
            
        # 2. DDL & DML Execution Prevention
        passed_cmd = True
        error_cmd = None
        if not getattr(request, "is_conversion", False):
            try:
                from ..guardrails.command_block import validate_command_injection
                validate_command_injection(request.logic)
                checked.append({"name": "DDL & DML Execution Prevention", "status": "Passed"})
            except HTTPException as he:
                passed_cmd = False
                error_cmd = he.detail
                checked.append({"name": "DDL & DML Execution Prevention", "status": "Failed", "message": error_cmd})
        else:
            checked.append({"name": "DDL & DML Execution Prevention", "status": "Passed", "message": "Bypassed for legacy code conversion."})
            
        # 3. Schema Access Protection
        passed_schema = True
        error_schema = None
        try:
            from ..guardrails.schema_protect import validate_schema_access
            validate_schema_access(request.userId, request.tables, request.columns)
            checked.append({"name": "Schema Access Protection", "status": "Passed"})
        except HTTPException as he:
            passed_schema = False
            error_schema = he.detail
            checked.append({"name": "Schema Access Protection", "status": "Failed", "message": error_schema})
            
        # 4. PII Redaction Guardrail
        checked.append({"name": "PII Redaction Guardrail", "status": "Passed", "message": "PII scanning active."})

        # 5. PII Access Protection
        passed_pii_access = True
        error_pii_access = None
        try:
            pii_cursor = db["piiForGuardrails"].find()
            for pii_doc in pii_cursor:
                param = pii_doc.get("piiParameter", "")
                reason = pii_doc.get("piiReason", "")
                pii_pass = pii_doc.get("piiPass", False)
                if param:
                    pattern = rf"\b{re.escape(param)}\b"
                    if re.search(pattern, request.logic, re.IGNORECASE):
                        if not pii_pass:
                            passed_pii_access = False
                            error_pii_access = f"Guardrail Violation: Access to sensitive PII parameter '{param}' is blocked (Reason: {reason})."
                            break
                    matched_col = None
                    for col in request.columns:
                        if re.search(pattern, col, re.IGNORECASE):
                            matched_col = col
                            break
                        if param in col:
                            matched_col = col
                            break
                    if matched_col and not pii_pass:
                        passed_pii_access = False
                        error_pii_access = f"Guardrail Violation: Access to sensitive PII column '{matched_col}' is blocked (Reason: {reason})."
                        break
            
            if passed_pii_access:
                checked.append({"name": "PII Access Protection", "status": "Passed", "message": "No sensitive PII requests detected."})
            else:
                checked.append({"name": "PII Access Protection", "status": "Failed", "message": error_pii_access})
        except Exception as pii_err:
            passed_pii_access = False
            error_pii_access = f"PII validation failed: {str(pii_err)}"
            checked.append({"name": "PII Access Protection", "status": "Failed", "message": error_pii_access})
            
        # 6. Domain Relevance Guardrail
        passed_relevance = True
        error_relevance = None
        try:
            from ..guardrails.relevance import validate_domain_relevance
            await validate_domain_relevance(request.logic, request.tables, request.columns, request.model)
            checked.append({"name": "Domain Relevance Guardrail", "status": "Passed"})
        except HTTPException as he:
            passed_relevance = False
            error_relevance = he.detail
            checked.append({"name": "Domain Relevance Guardrail", "status": "Failed", "message": error_relevance})

        passed = passed_pi and passed_bypass and passed_profanity and passed_cmd and passed_schema and passed_pii_access and passed_relevance
        first_error = error_pi or error_bypass or error_profanity or error_cmd or error_schema or error_pii_access or error_relevance
        
        return {
            "passed": passed,
            "error": first_error,
            "checked": checked
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

import io

def extract_printable_strings(bytes_data: bytes) -> str:
    import re
    # Find all contiguous sequences of printable ASCII characters of length 4 or more
    ascii_printables = re.findall(b'[\x20-\x7E\t\r\n]{4,}', bytes_data)
    decoded = []
    for chunk in ascii_printables:
        s = chunk.decode('ascii', errors='ignore').strip()
        if s:
            decoded.append(s)
    return "\n".join(decoded)

async def parse_uploaded_file(file_name: str, contents: bytes) -> str:
    ext = file_name.split('.')[-1].lower() if '.' in file_name else ''
    
    if ext in ['txt', 'xml', 'json']:
        return contents.decode('utf-8', errors='ignore')
        
    elif ext == 'pdf':
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(contents))
            text = ""
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
            return text.strip()
        except Exception as e:
            print(f"pypdf extraction failed: {e}. Falling back to printable ASCII extraction.")
            return extract_printable_strings(contents)
            
    elif ext in ['docx', 'doc']:
        if ext == 'docx':
            try:
                import docx
                doc = docx.Document(io.BytesIO(contents))
                text = "\n".join([p.text for p in doc.paragraphs])
                return text.strip()
            except Exception as e:
                print(f"python-docx extraction failed: {e}. Falling back to printable ASCII extraction.")
                return extract_printable_strings(contents)
        else: # .doc
            return extract_printable_strings(contents)
            
    else:
        try:
            return contents.decode('utf-8')
        except UnicodeDecodeError:
            return extract_printable_strings(contents)

@router.post("/generate/parse-legacy")
async def parse_legacy_file(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        file_name = file.filename or "legacy_code.txt"
        
        extracted_text = await parse_uploaded_file(file_name, contents)
        
        db = get_db()
        detected_tables, detected_columns = detect_tables_and_columns(extracted_text, db)
        
        return {
            "legacy_code": extracted_text,
            "detected_tables": detected_tables,
            "detected_columns": detected_columns
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse legacy file: {str(e)}")
